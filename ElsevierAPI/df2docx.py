from ElsevierAPI.pandas.panda_tricks import df
from ElsevierAPI.docx.docx_tools import set_repeat_table_header,add_hyperlink_into_run
from ElsevierAPI.docx.docx_tools import Document, add_caption
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from docx.shared import RGBColor
from ElsevierAPI.NCBI.pubmed import PUBMED_URL

PAGE_WIDTH = 8.5 #inches
TABLE_WIDTH = 7.5 # inches
TABLE_MARGIN = Inches((PAGE_WIDTH - TABLE_WIDTH)/2)


PG_COL = 'Click on # references to view practice guidelines'
SCILIT_COL = 'Click on # references to view relevant articles'

URLS = {'PMIDs or DOIs':PUBMED_URL,
        'References':PUBMED_URL,
        'Top References':PUBMED_URL,
        'Recent pubs':PUBMED_URL,
        'Top PMIDs':PUBMED_URL,
        'Recent PMIDs':PUBMED_URL,
        'PMID or DOI':PUBMED_URL,
        'Identifier':PUBMED_URL,
        PG_COL:PUBMED_URL,
        SCILIT_COL:PUBMED_URL,
        'Click on # references to view 5 most relevant articles in Pubmed':PUBMED_URL,
        'DOIs':'https://doi.org/',
        'Top DOIs':'https://doi.org/',
        'Recent DOIs':'https://doi.org/',
        'LOINC ID':'https://loinc.org/',
        'GV':'https://www.ncbi.nlm.nih.gov/snp/'
        }


def adjust_table_layout(table_layout:dict, to_width=TABLE_WIDTH):
    current_table_width = sum(list(table_layout.values()))
    scale_factor = to_width/current_table_width
    adjusted_layout = {k:v*scale_factor for k,v in table_layout.items()}
    return adjusted_layout


def df2table(document,clean_df:df,col2url=dict(),vertical_header=False,font_size=9,max_rows=0):
    my_df = df.copy_df(clean_df)
    my_df.fillna('', inplace=True)
    my_layout = my_df.inch_width_layout() if my_df.column2format else my_df.table_layout()
    my_layout = adjust_table_layout(my_layout)
    maximum_row_num = max_rows if max_rows else len(my_df)
    
    table_doc = document.add_table(rows=1, cols=len(my_df.columns))
    table_doc.autofit = False
    table_doc.allow_autofit = False

    current_section = document.sections[-1]
    current_section.left_margin = TABLE_MARGIN
    current_section.right_margin = TABLE_MARGIN  # need it to fit Inches(7.5) wide table
    
    hdr_cells = table_doc.rows[0].cells
    for i in range(0,len(my_df.columns)):
        hdr_cells[i].text = my_df.columns[i]
        hdr_cells[i].width = Inches(my_layout[i])
    
    row_couter = 0
    for idx in my_df.index:
        row = list(my_df.loc[idx])
        row_cells = table_doc.add_row().cells
        row_couter += 1

        for col_idx in range(0,len(row)):
            cell = row[col_idx]
            txt = str(cell) # cell in pandas can be int
            row_cells[col_idx].text = txt
            if txt:
                #cell_width = layout[col_idx]
                #font_size = cell_font_size(txt,cell_width,min_size=small_font)            
                row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(font_size)
                col_name = hdr_cells[col_idx].text
                if txt.isnumeric():
                    row_cells[col_idx].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                try:
                    url = col2url[col_name] # column must be listed in col2url to insert hyperink
                    if txt[0] == '=': # table has hyperlink
                        hyperlink,hyper_text = df.formula2hyperlink(txt)
                        row_cells[col_idx].text = hyper_text
                        
                        #hyper_text, hyperlink = add_hyperlink(row_cells[col_idx].paragraphs[0], url, hyper_text)
                    else:
                        hyperlink = url+txt
                    row_cells[col_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(6,69,173)
                    add_hyperlink_into_run(row_cells[col_idx].paragraphs[0], row_cells[col_idx].paragraphs[0].runs[0],hyperlink)
                    #hyperlink = add_hyperlink(row_cells[col_count].paragraphs[0], hyperlink, txt, '0645AD', True)
                except KeyError:
                    pass

            #every cell width must be set explicityly for layout to apply!!!!
            if str(row_cells[col_idx].text).isnumeric():  
                row_cells[col_idx].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                row_cells[col_idx].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            row_cells[col_idx].width = Inches(my_layout[col_idx])
        if row_couter >= maximum_row_num:
            break

    set_repeat_table_header(table_doc.rows[0],vertical_header)
    
# !!!!!!!!Table layout must FOLLOW table content fill in !!!!!!!!
    table_doc.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_doc.style = 'Grid Table 4'
    return

REF_TABLE = {0:0.9,1:5.5,2:0.9}
def add_ref_table(document:Document,tsv_in:str,table_name:str):
    """
    Input
    -----
    tsv_in format: ['Citation index', 'Biblio string', 'identifier']
    """
    ref_df = df.read_clean(tsv_in)
    ref_df.dropna(how='any',inplace=True)
    add_caption(document, table_name, level=3)
    df2table(document,ref_df,small_font=9,layout=REF_TABLE, col2url=URLS,vertical_header=False)


def df2ref_table(document:Document,ref_df:df,table_name:str):
    """
    Input
    -----
    tsv_in format: ['Citation index', 'Biblio string', 'identifier']
    """
    ref_df.dropna(how='any',inplace=True)
    add_caption(document, table_name, level=3)
    df2table(document,ref_df,font_size=9,layout=REF_TABLE, col2url=URLS,vertical_header=False)


