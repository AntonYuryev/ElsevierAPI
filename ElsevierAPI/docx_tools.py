from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.shared import OxmlElement,qn
from docx.shared import RGBColor
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.table import _Cell
import pypandoc 
from os.path import exists
import json

DECIMALcode = {'black':(0,0,0), 'white':(255,255,255),'lightgray':(119,136,153),'dark_blue':(0,0,255)}
HEXcode = {'black':'#000000', 'blue':'#0000EE'}
LEVEL2SPACER = {1:1.5, 2:0.4, 3:0.3}


def set_vertical_cell_direction(cell: _Cell, direction="tbRl"):
    # direction: tbRl -- top to bottom, btLr -- bottom to top
    assert direction in ("tbRl", "btLr")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)  # btLr tbRl
    tcPr.append(textDirection)
    return


def set_repeat_table_header(row,vertical_header=False):
    """ set repeat table row on every new page
    """
    col_count = 0
    got_vertical_header = False
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.bold = True
        font_color = DECIMALcode['white']
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(font_color[0],font_color[1],font_color[2])
        cell.paragraphs[0].alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
        #cell.paragraphs[0].paragraph_format.left_indent = Inches(0.1)

        tblCellProperties = cell._tc.get_or_add_tcPr()
        clShading = OxmlElement('w:shd')
        clShading.set(qn('w:fill'), HEXcode['black']) #Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
        tblCellProperties.append(clShading)

        #cell.paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.GRAY_25
        if vertical_header:
            if Inches(text_len(cell.text)) > cell.width:
                if cell.width < Inches(text_len('_partners_')):
                    set_vertical_cell_direction(cell)
                    got_vertical_header = True
        col_count += 1
    
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    if got_vertical_header:
        row.height = Inches(1)
    return row


def add_caption(doc:str or Document, table_name:str, legend='', level=3):
    document = Document(doc) if isinstance(doc, str) else doc
    try:
        heading_paragraph = document.add_heading(table_name, level=level)
    except KeyError:
        heading_paragraph = document.add_heading(table_name, level=6)

    try:
        spacing = Inches(LEVEL2SPACER[level])
    except KeyError:
        spacing = Inches(0.05)
    heading_paragraph.paragraph_format.space_before = spacing

    if legend:
        legend_paragraph = document.add_paragraph(legend)
    else:
        legend_paragraph = None
    if isinstance(doc, str): document.save(doc)

    return heading_paragraph, legend_paragraph

def text_len(text:str):
    """
    return text length in inches
    1 letter = 0.1 inches
    """
    return float(0.1*float(len(text)))

def cell_font_size(cell_text:str, cell_width:float, min_size=7, max_size=12):
    """
    cell_width must be in inches
    """
    cell_text_len = text_len(cell_text)
    proportinate_size = int(max_size*cell_width/cell_text_len)
    if proportinate_size <= min_size: return min_size
    elif proportinate_size >= max_size: return max_size
    return proportinate_size

def add_hyperlink_into_run(paragraph, run, url):
    runs = paragraph.runs
    for i in range(len(runs)):
        if runs[i].text == run.text:
            break

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    hyperlink.append(run._r)
    paragraph._p.insert(i+1,hyperlink)


def add_hyperlink(paragraph, url, text, color=HEXcode['blue'], underline=False):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = OxmlElement('w:color')
      c.set(qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = OxmlElement('w:u')
      u.set(qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def print_table_styles(document:Document):
    print('Current document has following table styles:')
    styles = [s for s in document.styles if s.type == WD_STYLE_TYPE.TABLE]
    for style in styles:
       print(style.name)
    print('Current document has following paragraph styles:')
    p_styles = [s for s in document.styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
    for style in p_styles:
       print(style.name)


def insert_paragraph_before(paragraph, before_text:str):
    return paragraph.insert_paragraph_before(before_text)


def fromHTML(html_fname:str):
    out_fname = html_fname[:-3] #assumes .xml extension
    if out_fname[-1] != '.':
        out_fname = out_fname[-1]
    from_formats, to_formats = pypandoc.get_pandoc_formats()
    print('pypandoc supports following input formats:')
    print(from_formats)
    print('pypandoc outputs formats:')
    print(to_formats)
    return pypandoc.convert_file(html_fname, 'docx', outputfile=out_fname+".docx")


def get_cell_size(cell_text:str, font_size=12):
    return len(cell_text)*font_size/12/8


def figure2doc(document, image_file, figure_title, legend, level=3):
    if exists(image_file):
        current_section = document.sections[-1]
        current_section.left_margin = Inches(0.5)
        current_section.right_margin = Inches(0.5)
        if figure_title:
            add_caption(document,figure_title,legend,level=level)
        document.add_picture(image_file, width=Inches(7.5))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        print ('Image from "%s" was added to the document' % image_file)
    else:
        print('Cannot find image file "%s"' % image_file)


def direct_insert(doc_dest, doc_src, figures_folder=''):
    #https://github.com/python-openxml/python-docx/blob/master/docx/text/parfmt.py
    style2indent={'Title':Inches(0),'Normal':Inches(1),'Heading 1':Inches(0),'Heading 2':Inches(0.25),'Heading 3':Inches(0.5),'Heading 4':Inches(0.75)}
    style2spacing={'Title':4,'Normal':0,'Heading 1':3,'Heading 2':2,'Heading 3':1,'Heading 4':0}
    if figures_folder:
        figure_attr = dict(json.load(open(figures_folder+'/figures.json','r')))
    
    for p in doc_src.paragraphs:
        inserted_p = doc_dest._body._body._insert_p(p._p)
        if p.text[:6] == 'Figure':
            figure_num_end = str(p.text).find(' ',8)
            figure_num = p.text[:figure_num_end]
            try:
                image_file, figure_title, legend = figure_attr[figure_num]
                figure2doc(doc_dest,figures_folder+'/'+image_file, figure_title, legend)
            except KeyError:
                print('Missing %s from %s folder' % (figure_num,figures_folder))

        my_style = p.style.name
        try:
            my_indent = style2indent[my_style]
            my_spacing = style2spacing[my_style]
        except KeyError:
            my_indent = Inches(1)
            my_spacing = -1

        inserted_p.get_or_add_pPr().ind_left = my_indent

        if p._p.get_or_add_pPr().numPr is not None:
            inserted_p.style = "ListNumber"
            my_level = p._p.get_or_add_pPr().numPr.ilvl.val
            inserted_p.get_or_add_pPr().ind_left = my_indent+ Inches(my_level*0.25)
            inserted_p.get_or_add_pPr().numPr.ilvl.val = my_level
            if my_spacing >= 0:
                inserted_p.get_or_add_pPr().spacing_after = my_spacing
            else:
                inserted_p.get_or_add_pPr().spacing_after = p._p.get_or_add_pPr().spacing_after

            inserted_p.get_or_add_pPr().spacing_before = 0
            inserted_p.get_or_add_pPr().spacing_line = 0

            if my_level > 0:
                inserted_p.get_or_add_pPr().first_line_indent = Inches(0.25/my_level)
            else:
                inserted_p.get_or_add_pPr().first_line_indent = Inches(0.3)
            
            inserted_p.get_or_add_pPr().spacing_lineRule = p._p.get_or_add_pPr().spacing_lineRule
            inserted_p.get_or_add_pPr().jc_val = p._p.get_or_add_pPr().jc_val


def add_figures(document, section_name:str, folder='figures'):
    fname = folder+'/figures.json'
    try:
        figure_attr = dict(json.load(open(fname,'r')))
        add_caption(document,section_name, level=1)
        for attrs in figure_attr.values():
            figure2doc(document, folder+'/'+attrs[0], attrs[1], attrs[2], level=2)
    except FileNotFoundError:
        print('Cannot find "%s" file\nNo section "%s" was created' % (fname,section_name))
